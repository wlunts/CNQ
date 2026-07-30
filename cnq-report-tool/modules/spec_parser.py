# -*- coding: utf-8 -*-
"""
验货资料解析器 - Spec Parser
=============================
从验货资料（规格书/订单要求/技术文档）中解析规格标准值。

支持格式：
- Excel (.xlsx/.xls) — 规格表、BOM 表
- Word (.docx)      — 技术规格书
- PDF               — 产品规格文档
- CSV/Text          — 简单数据文件
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Any

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class SpecParser:
    """
    验货资料解析器
    
    从各种格式的验货资料中提取规格标准值，
    用于后续与实际测量数据进行比对。
    """
    
    # 规格字段关键词（用于识别资料中的相关列/段）
    SPEC_FIELD_KEYWORDS = {
        "product_name":    ["产品名称", "品名", "Product Name", "Description", "Item", "物料名称"],
        "product_dims":    ["产品尺寸", "Product Size", "Item Size", "Dimension", "规格", "Size"],
        "product_length":  ["产品长", "Length", "L", "长"],
        "product_width":   ["产品宽", "Width", "W", "宽"],
        "product_height":  ["产品高", "Height", "H", "高", "Thickness", "厚度"],
        "carton_dims":     ["外箱尺寸", "Carton Size", "箱规", "Outer Dimension"],
        "carton_length":   ["外箱长", "Carton L"],
        "carton_width":    ["外箱宽", "Carton W"],
        "carton_height":   ["外箱高", "Carton H"],
        "gross_weight":    ["毛重", "Gross Weight", "G.W.", "Gross WT"],
        "net_weight":      ["净重", "Net Weight", "N.W.", "Net WT"],
        "qty_per_carton":  ["装箱数", "Qty/Ctn", "PCS/CTN", "入数", "Pack Qty"],
        "material":        ["材质", "Material", "材料"],
        "color":           ["颜色", "Color", "Colour"],
        "tolerance_dims":  ["尺寸公差", "Tolerance", "允差", "Dims Tolerance"],
        "tolerance_weight":["重量公差", "Weight Tolerance"],
    }
    
    DIMS_PATTERN = re.compile(
        r'(\d+\.?\d*)\s*[xX×\*]\s*(\d+\.?\d*)\s*[xX×\*]\s*(\d+\.?\d*)'
    )
    
    NUMERIC_PATTERN = re.compile(r'(\d+\.?\d*)')
    
    def __init__(self):
        pass
    
    # ── 主入口 ────────────────────────────────────────────
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析验货资料文件
        
        Args:
            file_path: 验货资料文件路径
        
        Returns:
            {field_name: value} 规格标准值
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        print(f"\n📋 解析验货资料: {path.name}")
        
        if suffix in ('.xlsx', '.xlsm', '.xls'):
            return self._parse_excel(file_path)
        elif suffix == '.docx':
            return self._parse_docx(file_path)
        elif suffix == '.pdf':
            return self._parse_pdf(file_path)
        elif suffix in ('.csv', '.txt'):
            return self._parse_text(file_path)
        else:
            print(f"  ⚠ 不支持的文件格式: {suffix}，尝试作为文本解析")
            return self._parse_text(file_path)
    
    # ── Excel 解析 ────────────────────────────────────────
    
    def _parse_excel(self, file_path: str) -> Dict[str, Any]:
        """解析 Excel 规格表"""
        if not HAS_OPENPYXL:
            raise ImportError("请安装 openpyxl: pip install openpyxl")
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        specs = {}
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            print(f"  Sheet: {sheet_name} ({ws.max_row}行 x {ws.max_column}列)")
            
            # 策略1：查找 key-value 对（label 在 A 列，value 在 B 列）
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, 
                                     min_col=1, max_col=min(ws.max_column, 3),
                                     values_only=False):
                label_cell = row[0]
                label = str(label_cell.value).strip() if label_cell.value else ""
                
                if not label:
                    continue
                
                matched_field = self._match_field(label)
                if matched_field and len(row) >= 2:
                    value_cell = row[1]
                    value = self._clean_value(value_cell.value)
                    if value is not None:
                        specs[matched_field] = value
                
                # 有些表格 value 在 C 列
                if matched_field and len(row) >= 3 and matched_field not in specs:
                    value_cell = row[2]
                    value = self._clean_value(value_cell.value)
                    if value is not None:
                        specs[matched_field] = value
            
            # 策略2：解析首行为表头的表格
            if ws.max_row > 1:
                specs.update(self._parse_header_table(ws))
        
        wb.close()
        specs = self._post_process(specs)
        self._print_specs(specs)
        return specs
    
    def _parse_header_table(self, ws) -> Dict[str, Any]:
        """解析首行为表头的表格，提取 key-value"""
        specs = {}
        headers = []
        for cell in ws[1]:
            headers.append(str(cell.value).strip() if cell.value else "")
        
        # 在 header 中查找规格字段
        for col_idx, header in enumerate(headers):
            matched = self._match_field(header)
            if matched:
                # 取第二行的值
                if ws.max_row >= 2:
                    val_cell = ws.cell(row=2, column=col_idx + 1)
                    value = self._clean_value(val_cell.value)
                    if value is not None:
                        specs[matched] = value
        
        return specs
    
    # ── Word 解析 ─────────────────────────────────────────
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 规格书"""
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        specs = {}
        full_text_parts = []
        
        # 解析表格
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                if len(cells_text) >= 2:
                    for i, cell_text in enumerate(cells_text):
                        matched = self._match_field(cell_text)
                        if matched:
                            # 取相邻右列或下一行的值
                            if i + 1 < len(cells_text):
                                value = self._clean_value(cells_text[i + 1])
                                if value is not None:
                                    specs[matched] = value
        
        # 解析段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_parts.append(text)
        
        # 从连续文本中提取
        full_text = "\n".join(full_text_parts)
        specs.update(self._extract_from_text(full_text))
        
        specs = self._post_process(specs)
        self._print_specs(specs)
        return specs
    
    # ── PDF 解析 ──────────────────────────────────────────
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 规格文档"""
        if not HAS_PDFPLUMBER:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")
        
        specs = {}
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    for row in table:
                        if not row:
                            continue
                        clean_row = [str(c).strip() if c else "" for c in row]
                        for i, cell_text in enumerate(clean_row):
                            matched = self._match_field(cell_text)
                            if matched and i + 1 < len(clean_row):
                                value = self._clean_value(clean_row[i + 1])
                                if value is not None:
                                    specs[matched] = value
                
                # 提取文本
                text = page.extract_text()
                if text:
                    specs.update(self._extract_from_text(text))
        
        specs = self._post_process(specs)
        self._print_specs(specs)
        return specs
    
    # ── 纯文本解析 ────────────────────────────────────────
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析纯文本 / CSV"""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        
        specs = self._extract_from_text(text)
        specs = self._post_process(specs)
        self._print_specs(specs)
        return specs
    
    def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """从自由文本中提取规格数据"""
        specs = {}
        lines = text.split("\n")
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 尝试 key: value 模式
            for sep in [":", "：", "=", "\t", "  "]:
                if sep in line:
                    parts = line.split(sep, 1)
                    if len(parts) == 2:
                        key_part = parts[0].strip()
                        val_part = parts[1].strip()
                        
                        matched = self._match_field(key_part)
                        if matched:
                            value = self._clean_value(val_part)
                            if value is not None:
                                specs[matched] = value
                                break
        
        return specs
    
    # ── 辅助方法 ──────────────────────────────────────────
    
    def _match_field(self, text: str) -> Optional[str]:
        """识别文本对应的规格字段名"""
        text_lower = text.lower().strip()
        
        # 去除多余空格和特殊字符
        text_clean = re.sub(r'[\s\-_]+', '', text_lower)
        
        for field, keywords in self.SPEC_FIELD_KEYWORDS.items():
            for kw in keywords:
                kw_clean = re.sub(r'[\s\-_]+', '', kw.lower())
                if kw_clean in text_clean or text_clean in kw_clean:
                    return field
        
        return None
    
    def _clean_value(self, value: Any) -> Optional[Any]:
        """清理和标准化值"""
        if value is None:
            return None
        
        if isinstance(value, (int, float)):
            return value
        
        text = str(value).strip()
        if not text or text.lower() in ("n/a", "na", "-", "--", "none", "null"):
            return None
        
        # 尝试从带标注的值中提取数字
        # "60 x 40 x 30 cm" → 保留原样以供后续解析
        dims_match = self.DIMS_PATTERN.search(text)
        if dims_match:
            l, w, h = dims_match.groups()
            return {
                "raw": text,
                "length": float(l),
                "width": float(w),
                "height": float(h),
            }
        
        # 纯数字（带可能的单位）
        num_match = self.NUMERIC_PATTERN.search(text)
        if num_match and len(num_match.group()) == len(re.sub(r'[a-zA-Z\s]', '', text)):
            try:
                return float(num_match.group(1))
            except ValueError:
                pass
        
        return text
    
    def _post_process(self, specs: Dict[str, Any]) -> Dict[str, Any]:
        """后处理：合并相关字段，标准化格式"""
        processed = dict(specs)
        
        # 合并 carton_length/width/height → carton_dims
        carton_parts = {}
        for key in ["carton_length", "carton_width", "carton_height"]:
            if key in processed:
                carton_parts[key] = processed.pop(key)
                carton_parts[key.replace("carton_", "")] = carton_parts[key]
        
        if len(carton_parts) >= 3:
            processed["carton_dims"] = {
                "length": carton_parts.get("length", carton_parts.get("carton_length", 0)),
                "width": carton_parts.get("width", carton_parts.get("carton_width", 0)),
                "height": carton_parts.get("height", carton_parts.get("carton_height", 0)),
            }
        elif carton_parts:
            # 部分数据，保留原样
            processed.update(carton_parts)
        
        # 合并 product_length/width/height → product_dims
        prod_parts = {}
        for key in ["product_length", "product_width", "product_height"]:
            if key in processed:
                val = processed.pop(key)
                if isinstance(val, (int, float)):
                    prod_parts[key.replace("product_", "")] = val
                elif isinstance(val, dict):
                    return val  # 已经是完整尺寸
        
        if len(prod_parts) >= 3:
            processed["product_dims"] = {
                "length": prod_parts["length"],
                "width": prod_parts["width"],
                "height": prod_parts["height"],
            }
        elif prod_parts:
            processed.update(prod_parts)
        
        # 展开尺寸 dict 中已解析的值
        for dims_key in ["product_dims", "carton_dims"]:
            if dims_key in processed:
                val = processed[dims_key]
                if isinstance(val, str):
                    match = self.DIMS_PATTERN.search(val)
                    if match:
                        processed[dims_key] = {
                            "length": float(match.group(1)),
                            "width": float(match.group(2)),
                            "height": float(match.group(3)),
                        }
        
        return processed
    
    def _print_specs(self, specs: Dict[str, Any]):
        """打印解析结果摘要"""
        if not specs:
            print("  ⚠ 未从资料中提取到规格数据")
            return
        
        print(f"  提取到 {len(specs)} 项规格:")
        for key, value in specs.items():
            if isinstance(value, dict):
                print(f"    {key}: {value}")
            else:
                print(f"    {key}: {value}")


# ── 快速入口 ─────────────────────────────────────────────

def parse_specs(file_path: str) -> Dict[str, Any]:
    """快速解析验货资料"""
    parser = SpecParser()
    return parser.parse(file_path)
