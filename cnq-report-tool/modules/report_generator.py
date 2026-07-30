# -*- coding: utf-8 -*-
"""
报告生成模块 - Report Generator
===============================
综合所有数据，生成最终的验货报告。

工作流：
1. 接收 PhotoAnalyzer 提取的实测数据
2. 接收 SpecParser 解析的规格标准
3. 接收 DataComparator 的比对结果
4. 参照参考报告的格式
5. 填充模板并输出最终 .docx 报告
"""

import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple

from config import DEFAULT_OUTPUT_DIR

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportGenerator:
    """
    验货报告生成器
    
    将所有模块的输出合并，生成最终报告。
    支持两种模式：
    1. 模板填充模式：基于空白模板填入数据
    2. 参考克隆模式：复制参考报告结构并替换数据
    """
    
    def __init__(self, output_dir: str = None):
        self.output_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
    
    # ── 模式1：基于空白模板生成 ──────────────────────────
    
    def generate_from_template(
        self,
        template_path: str,
        extracted_data: Dict[str, Any],
        specs_data: Dict[str, Any],
        comparison_result: Dict[str, Any],
        reference_path: str = None,
        output_filename: str = None,
    ) -> str:
        """
        从空白模板生成报告
        
        Args:
            template_path: 空白报告模板 (.docx)
            extracted_data: 照片提取的实测数据
            specs_data: 验货资料解析的规格
            comparison_result: 比对结果
            reference_path: 参考报告（用于格式学习，可选）
            output_filename: 输出文件名（可选）
        
        Returns:
            输出文件路径
        """
        from .template_handler import TemplateHandler
        
        print("\n" + "█" * 60)
        print("  📝 正在生成验货报告...")
        print("█" * 60)
        
        handler = TemplateHandler()
        
        # 1. 分析参考报告（如果提供）
        ref_info = None
        if reference_path and Path(reference_path).exists():
            ref_info = handler.analyze_reference(reference_path)
        
        # 2. 解析模板
        doc, structure = handler.parse_template(template_path)
        
        # 3. 准备填充数据
        report_data = handler.prepare_report_data(
            extracted=extracted_data,
            specs=specs_data,
            comparisons=comparison_result.get("comparisons", []),
        )
        
        # 添加比对结果
        report_data["Overall_Result"] = comparison_result.get("overall_result", "")
        report_data["Passed_Count"] = str(comparison_result.get("passed_count", 0))
        report_data["Failed_Count"] = str(comparison_result.get("failed_count", 0))
        report_data["Remarks"] = comparison_result.get("remarks", "")
        report_data["Report_Date"] = datetime.now().strftime("%Y-%m-%d")
        
        # 4. 填充模板
        doc = handler.fill_template(doc, report_data, structure)
        
        # 5. 保存
        if output_filename is None:
            po = extracted_data.get("po_number", "unknown")
            date_str = datetime.now().strftime("%Y%m%d")
            output_filename = f"Inspection_Report_{po}_{date_str}.docx"
        
        output_path = self.output_dir / output_filename
        handler.save_document(doc, str(output_path))
        
        # 6. 打印摘要
        self._print_report_summary(extracted_data, comparison_result, str(output_path))
        
        return str(output_path)
    
    # ── 模式2：基于参考报告克隆 ──────────────────────────
    
    def generate_from_reference(
        self,
        reference_path: str,
        extracted_data: Dict[str, Any],
        specs_data: Dict[str, Any],
        comparison_result: Dict[str, Any],
        output_filename: str = None,
    ) -> str:
        """
        基于参考报告生成新报告（克隆格式并替换数据）
        
        这种模式更适合已有完整参考报告的场景，
        直接复制参考报告的样式和结构。
        
        Args:
            reference_path: 参考完好的报告 (.docx)
            extracted_data: 照片提取的实测数据
            specs_data: 验货资料解析的规格
            comparison_result: 比对结果
            output_filename: 输出文件名
        
        Returns:
            输出文件路径
        """
        print("\n" + "█" * 60)
        print("  📝 基于参考报告生成新报告...")
        print("█" * 60)
        
        doc = Document(reference_path)
        
        # 构建替换映射
        replacements = self._build_replacements(extracted_data, specs_data, comparison_result)
        
        # 替换段落文本
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)
        
        # 替换表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)
        
        # 保存
        if output_filename is None:
            po = extracted_data.get("po_number", "unknown")
            date_str = datetime.now().strftime("%Y%m%d")
            output_filename = f"Inspection_Report_{po}_{date_str}.docx"
        
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        print(f"\n✅ 报告已保存: {output_path}")
        self._print_report_summary(extracted_data, comparison_result, str(output_path))
        
        return str(output_path)
    
    # ── 替换逻辑 ──────────────────────────────────────────
    
    def _build_replacements(self, extracted: Dict, specs: Dict,
                            comparison: Dict) -> Dict[str, str]:
        """构建参考报告中的文本替换映射"""
        from .template_handler import TemplateHandler
        handler = TemplateHandler()
        
        report_data = handler.prepare_report_data(
            extracted=extracted,
            specs=specs,
            comparisons=comparison.get("comparisons", []),
        )
        
        # 添加比对汇总
        report_data["Overall_Result"] = comparison.get("overall_result", "")
        report_data["Passed_Count"] = str(comparison.get("passed_count", 0))
        report_data["Failed_Count"] = str(comparison.get("failed_count", 0))
        report_data["Report_Date"] = datetime.now().strftime("%Y-%m-%d")
        
        # 把 dict 值也转为字符串
        for key, val in report_data.items():
            if isinstance(val, (int, float)):
                report_data[key] = str(val)
            elif isinstance(val, dict):
                report_data[key] = str(val)
        
        return report_data
    
    def _replace_in_paragraph(self, para, replacements: Dict[str, str]):
        """在段落中执行文本替换"""
        for key, value in replacements.items():
            if key in para.text or key.lower() in para.text.lower():
                # 替换 key: value 模式（参考报告中的标签）
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                    elif key.lower() in run.text.lower():
                        # 大小写不敏感替换
                        run.text = self._case_insensitive_replace(run.text, key, str(value))
    
    def _case_insensitive_replace(self, text: str, old: str, new: str) -> str:
        """大小写不敏感的文本替换"""
        import re
        return re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
    
    # ── 模式3：纯数据输出（JSON + 摘要） ──────────────────
    
    def generate_data_summary(
        self,
        extracted_data: Dict[str, Any],
        specs_data: Dict[str, Any],
        comparison_result: Dict[str, Any],
        output_filename: str = None,
    ) -> str:
        """
        生成数据摘要文件（JSON 格式）
        
        用于无模板时的数据输出，方便手动导入报告。
        """
        import json
        
        summary = {
            "report_metadata": {
                "generated_at": datetime.now().isoformat(),
                "tool_version": "1.0.0",
            },
            "extracted_data": extracted_data,
            "specifications": specs_data,
            "comparison": {
                "overall_result": comparison_result.get("overall_result"),
                "passed_count": comparison_result.get("passed_count"),
                "failed_count": comparison_result.get("failed_count"),
                "skipped_count": comparison_result.get("skipped_count"),
                "remarks": comparison_result.get("remarks"),
                "items": comparison_result.get("comparisons", []),
            },
        }
        
        if output_filename is None:
            po = extracted_data.get("po_number", "unknown")
            output_filename = f"report_data_{po}.json"
        
        output_path = self.output_dir / output_filename
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(summary, f, ensure_ascii=False, indent=2, default=str)
        
        print(f"\n📄 数据摘要已导出: {output_path}")
        return str(output_path)
    
    # ── 打印摘要 ──────────────────────────────────────────
    
    def _print_report_summary(self, extracted: Dict, comparison: Dict, output_path: str):
        """打印报告生成摘要"""
        print("\n" + "─" * 50)
        print("  📋 报告生成摘要")
        print("─" * 50)
        
        if extracted.get("po_number"):
            print(f"  PO 号: {extracted['po_number']}")
        if extracted.get("product_name"):
            print(f"  产品: {extracted['product_name']}")
        if extracted.get("carton_dims"):
            d = extracted["carton_dims"]
            print(f"  箱规: {d['length']} x {d['width']} x {d['height']} cm")
        
        print(f"  比对结果: {comparison.get('overall_result', 'N/A')}")
        print(f"  通过/不合格/跳过: "
              f"{comparison.get('passed_count', 0)}/"
              f"{comparison.get('failed_count', 0)}/"
              f"{comparison.get('skipped_count', 0)}")
        print(f"  输出路径: {output_path}")
        print("─" * 50)


# ── 快捷入口 ─────────────────────────────────────────────

def generate_report(
    template_path: str,
    extracted_data: Dict,
    specs_data: Dict,
    comparison_result: Dict,
    reference_path: str = None,
    output_filename: str = None,
) -> str:
    """一键生成报告"""
    gen = ReportGenerator()
    
    if Path(template_path).exists():
        # 有模板 → 模板填充模式
        return gen.generate_from_template(
            template_path, extracted_data, specs_data,
            comparison_result, reference_path, output_filename,
        )
    elif reference_path and Path(reference_path).exists():
        # 无模板但有参考报告 → 克隆模式
        return gen.generate_from_reference(
            reference_path, extracted_data, specs_data,
            comparison_result, output_filename,
        )
    else:
        # 都没有 → 输出 JSON 数据摘要
        print("\n⚠ 未提供模板或参考报告，生成 JSON 数据摘要...")
        return gen.generate_data_summary(
            extracted_data, specs_data, comparison_result, output_filename,
        )
