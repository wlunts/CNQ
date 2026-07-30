# -*- coding: utf-8 -*-
"""
报告模板处理器 - Template Handler
==================================
读取、解析和填充 Word (.docx) 验货报告模板。

核心功能：
1. 读取空白模板 → 识别表格结构和字段位置
2. 读取参考报告 → 学习格式（表格样式、字体、布局）
3. 将数据填入模板对应位置
"""

import re
import copy
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from config import REPORT_TABLE_HEADERS, PLACEHOLDER_PATTERN


class TemplateStructure:
    """模板结构描述"""
    
    def __init__(self):
        self.tables: List[Dict] = []          # 表格列表，每项 {table_index, rows, cols, headers, ...}
        self.placeholders: Dict[str, List] = {}  # 占位符位置 {field: [(para_idx, run_idx, table_idx), ...]}
        self.section_indices: Dict[str, int] = {}  # 段落/表格索引映射
        
    def to_dict(self) -> dict:
        return {
            "tables": self.tables,
            "placeholders": {k: [(p[0], p[1], p[2]) for p in v] for k, v in self.placeholders.items()},
            "sections": self.section_indices,
        }


class TemplateHandler:
    """
    验货报告模板处理器
    
    支持：
    - 解析模板结构
    - 识别 {{占位符}} 和 【占位符】
    - 表格数据定位与填充
    - 从参考报告学习格式
    """
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
    
    # ── 模板解析 ──────────────────────────────────────────
    
    def parse_template(self, template_path: str) -> Tuple[Document, TemplateStructure]:
        """
        解析空白模板，识别结构和占位符
        
        Returns:
            (Document 对象, TemplateStructure 结构)
        """
        doc = Document(template_path)
        structure = TemplateStructure()
        
        print(f"\n📄 解析模板: {Path(template_path).name}")
        
        # 1. 解析段落中的占位符
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            placeholders = PLACEHOLDER_PATTERN.findall(text)
            for ph in placeholders:
                # ph 是元组，取第一个非空值
                key = next((p for p in ph if p), "")
                if key:
                    key = key.strip()
                    if key not in structure.placeholders:
                        structure.placeholders[key] = []
                    structure.placeholders[key].append((i, -1, -1))  # (para_idx, run_idx=-1, table_idx=-1)
        
        # 2. 解析表格
        print(f"  发现 {len(doc.tables)} 个表格")
        for t_idx, table in enumerate(doc.tables):
            headers = []
            # 尝试从第一行获取表头
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            table_info = {
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "headers": headers,
            }
            
            # 识别表格类型（属于报告的哪个部分）
            section = self._identify_table_section(headers, table)
            table_info["section"] = section
            structure.section_indices[section] = t_idx
            
            # 解析表格中的占位符
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        placeholders = PLACEHOLDER_PATTERN.findall(para.text)
                        for ph in placeholders:
                            key = next((p for p in ph if p), "")
                            if key:
                                key = key.strip()
                                if key not in structure.placeholders:
                                    structure.placeholders[key] = []
                                structure.placeholders[key].append((-1, -1, (t_idx, r_idx, c_idx)))
            
            structure.tables.append(table_info)
        
        print(f"  找到 {len(structure.placeholders)} 个占位符字段")
        for key, locs in structure.placeholders.items():
            print(f"    - {key}: {len(locs)} 处")
        
        return doc, structure
    
    def _identify_table_section(self, headers: List[str], table) -> str:
        """根据表头关键词识别表格类型"""
        header_text = " ".join(headers).lower()
        
        for section, keywords in REPORT_TABLE_HEADERS.items():
            for kw in keywords:
                if kw.lower() in header_text:
                    return section
        
        # 检查表体内容
        full_text = ""
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text.lower() + " "
        
        for section, keywords in REPORT_TABLE_HEADERS.items():
            for kw in keywords:
                if kw.lower() in full_text:
                    return section
        
        return "unknown"
    
    # ── 参考报告分析 ──────────────────────────────────────
    
    def analyze_reference(self, reference_path: str) -> Dict[str, Any]:
        """
        分析参考报告，提取格式信息
        
        Returns:
            {tables: [...], fonts: {...}, styles: {...}, sections: {...}}
        """
        doc = Document(reference_path)
        ref_info = {
            "file": Path(reference_path).name,
            "tables": [],
            "fonts": {},
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "sections": {},
        }
        
        print(f"\n📋 分析参考报告: {ref_info['file']}")
        
        # 提取表格结构
        for t_idx, table in enumerate(doc.tables):
            headers = []
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            section = self._identify_table_section(headers, table)
            
            # 提取表格样式
            style_info = {
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "headers": headers,
                "section": section,
            }
            ref_info["tables"].append(style_info)
            ref_info["sections"][section] = t_idx
        
        # 提取字体信息
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            if first_para.runs:
                run = first_para.runs[0]
                ref_info["fonts"]["default_size"] = run.font.size
                ref_info["fonts"]["default_name"] = run.font.name
        
        print(f"  段落: {ref_info['paragraph_count']}")
        print(f"  表格: {ref_info['table_count']}")
        for t in ref_info["tables"]:
            print(f"    [{t['section']}] {t['rows']}行 x {t['cols']}列: {t['headers'][:3]}...")
        
        return ref_info
    
    # ── 数据填充 ──────────────────────────────────────────
    
    def fill_template(self, doc: Document, data: Dict[str, Any],
                      structure: TemplateStructure = None) -> Document:
        """
        将数据填入模板
        
        Args:
            doc: python-docx Document 对象
            data: 要填充的数据，key 对应占位符名
            structure: 模板结构（可选，用于精确定位）
        
        Returns:
            修改后的 Document
        """
        # 构建查找映射：data key 的多种可能写法
        key_map = self._build_key_map(data)
        
        # 1. 填充段落中的占位符
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            self._fill_paragraph_placeholders(para, data, key_map)
        
        # 2. 填充表格中的占位符
        for table in doc.tables:
            self._fill_table_data(table, data, key_map, structure)
        
        return doc
    
    def _build_key_map(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """构建 key → value 的多种映射，处理各种命名变体"""
        key_map = {}
        for key, value in data.items():
            key_map[key] = value
            key_map[key.lower()] = value
            key_map[key.upper()] = value
            key_map[key.replace("_", " ")] = value
            key_map[key.replace("_", " ").title()] = value
            key_map[key.replace("_", "").lower()] = value
        
        # 展开嵌套 dict
        for key, value in data.items():
            if isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    combined = f"{key}_{sub_key}"
                    key_map[combined] = sub_val
                    key_map[combined.lower()] = sub_val
        
        return key_map
    
    def _fill_paragraph_placeholders(self, para, data: Dict, key_map: Dict):
        """替换段落中的 {{placeholder}} 占位符"""
        runs = para.runs
        
        # 合并所有 runs 的文本
        full_text = "".join(r.text for r in runs)
        
        # 查找并替换占位符
        modified = False
        new_text = full_text
        
        for match in PLACEHOLDER_PATTERN.finditer(full_text):
            ph_key = next((g for g in match.groups() if g), "")
            if not ph_key:
                continue
            ph_key = ph_key.strip()
            
            # 在 key_map 中查找匹配
            replacement = self._find_value(ph_key, key_map)
            if replacement is not None:
                new_text = new_text.replace(match.group(0), str(replacement))
                modified = True
        
        if modified:
            # 清除所有 runs 并设置新文本
            for run in runs:
                run.text = ""
            if runs:
                runs[0].text = new_text
    
    def _fill_table_data(self, table, data: Dict, key_map: Dict, structure=None):
        """填充表格数据"""
        for row in table.rows:
            for cell in row.cells:
                # 检查单元格中的占位符
                cell_text = cell.text
                
                for match in PLACEHOLDER_PATTERN.finditer(cell_text):
                    ph_key = next((g for g in match.groups() if g), "")
                    if not ph_key:
                        continue
                    ph_key = ph_key.strip()
                    replacement = self._find_value(ph_key, key_map)
                    
                    if replacement is not None:
                        # 替换单元格中所有段落的占位符
                        for para in cell.paragraphs:
                            para_text = para.text
                            if match.group(0) in para_text:
                                new_text = para_text.replace(match.group(0), str(replacement))
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = new_text
    
    def _find_value(self, ph_key: str, key_map: Dict) -> Optional[Any]:
        """在 key_map 中查找占位符对应的值"""
        # 直接匹配
        if ph_key in key_map:
            return key_map[ph_key]
        
        # 大小写不敏感
        ph_lower = ph_key.lower().strip()
        for key, val in key_map.items():
            if key.lower().strip() == ph_lower:
                return val
        
        # 模糊匹配（占位符是 key 的子串，反之亦然）
        for key, val in key_map.items():
            if ph_lower in key.lower() or key.lower() in ph_lower:
                return val
        
        return None
    
    # ── 格式化辅助 ────────────────────────────────────────
    
    def format_dimensions(self, dims: Dict[str, float], unit: str = "cm") -> str:
        """格式化尺寸为字符串"""
        if not dims:
            return ""
        l = dims.get("length", dims.get("l", 0))
        w = dims.get("width", dims.get("w", 0))
        h = dims.get("height", dims.get("h", 0))
        return f"{l} x {w} x {h} {unit}"
    
    def format_weight(self, value: Optional[float], unit: str = "kg") -> str:
        """格式化重量"""
        if value is None:
            return ""
        return f"{value:.2f} {unit}"
    
    def prepare_report_data(self, extracted: Dict, specs: Dict = None,
                            comparisons: List[Dict] = None) -> Dict[str, Any]:
        """
        准备填充报告的完整数据集
        
        将原始提取数据转换为报告模板需要的格式
        """
        data = {}
        
        # 产品基本信息
        data["Product_Name"] = extracted.get("product_name", "")
        data["PO_Number"] = extracted.get("po_number", "")
        data["Material"] = extracted.get("material", "")
        data["Color"] = extracted.get("color", "")
        data["Barcode_SKU"] = extracted.get("barcode_sku", "")
        data["MFG_Date"] = extracted.get("mfg_date", "")
        data["Batch_Lot"] = extracted.get("batch_lot", "")
        
        # 外箱信息
        carton = extracted.get("carton_dims", {})
        data["Carton_Dimensions"] = self.format_dimensions(carton)
        data["Carton_Length"] = str(carton.get("length", ""))
        data["Carton_Width"] = str(carton.get("width", ""))
        data["Carton_Height"] = str(carton.get("height", ""))
        
        # 产品尺寸
        prod = extracted.get("product_dims", {})
        data["Product_Dimensions"] = self.format_dimensions(prod)
        data["Product_Length"] = str(prod.get("length", ""))
        data["Product_Width"] = str(prod.get("width", ""))
        data["Product_Height"] = str(prod.get("height", ""))
        
        # 重量
        wt_unit = extracted.get("weight_unit", "kg")
        data["Gross_Weight"] = self.format_weight(extracted.get("gross_weight"), wt_unit)
        data["Net_Weight"] = self.format_weight(extracted.get("net_weight"), wt_unit)
        data["Weight_Unit"] = wt_unit
        
        # 数量
        data["Qty_Per_Carton"] = str(extracted.get("qty_per_carton", ""))
        data["Total_Quantity"] = str(extracted.get("total_quantity", ""))
        
        # 规格参考值（如果有）
        if specs:
            spec_carton = specs.get("carton_dims", {})
            data["Spec_Carton_Dims"] = self.format_dimensions(spec_carton)
            spec_prod = specs.get("product_dims", {})
            data["Spec_Product_Dims"] = self.format_dimensions(spec_prod)
            data["Spec_Gross_Weight"] = self.format_weight(
                specs.get("gross_weight"), specs.get("weight_unit", "kg")
            )
            data["Spec_Net_Weight"] = self.format_weight(
                specs.get("net_weight"), specs.get("weight_unit", "kg")
            )
        
        # 不合格项备注
        if comparisons:
            remarks_lines = []
            for i, comp in enumerate(comparisons, 1):
                if not comp.get("passed", True):
                    remarks_lines.append(
                        f"{i}. [{comp.get('field', '')}] "
                        f"实测: {comp.get('actual', '')} | "
                        f"标准: {comp.get('expected', '')} → "
                        f"{comp.get('remark', '不合格')}"
                    )
            data["Remarks"] = "\n".join(remarks_lines) if remarks_lines else "无"
            data["Non_Conformity_Count"] = str(sum(1 for c in comparisons if not c.get("passed", True)))
        
        # 中文别名
        data["产品名称"] = data["Product_Name"]
        data["PO号"] = data["PO_Number"]
        data["材质"] = data["Material"]
        data["箱规"] = data["Carton_Dimensions"]
        data["毛重"] = data["Gross_Weight"]
        data["净重"] = data["Net_Weight"]
        data["装箱数"] = data["Qty_Per_Carton"]
        data["备注"] = data.get("Remarks", "")
        data["不合格项数"] = data.get("Non_Conformity_Count", "0")
        
        return data
    
    # ── 保存 ──────────────────────────────────────────────
    
    def save_document(self, doc: Document, output_path: str):
        """保存修改后的文档"""
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print(f"\n✅ 报告已保存: {out}")


# ── 快速入口 ─────────────────────────────────────────────

def process_template(template_path: str, data: Dict[str, Any],
                     output_path: str = None) -> str:
    """
    快速处理：解析模板 → 填充数据 → 保存
    
    Args:
        template_path: 模板文件路径
        data: 填充数据
        output_path: 输出路径（默认: 模板名_filled.docx）
    
    Returns:
        输出文件路径
    """
    handler = TemplateHandler()
    doc, structure = handler.parse_template(template_path)
    doc = handler.fill_template(doc, data, structure)
    
    if output_path is None:
        p = Path(template_path)
        output_path = str(p.parent / f"{p.stem}_filled{p.suffix}")
    
    handler.save_document(doc, output_path)
    return output_path
